"""python-docx ile APA'ya yakın Word çıktısı."""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.shared import Cm, Pt

from app.literature.apa import format_reference_runs
from app.models import CleaningReport, DiscoveryReport, Finding, Manuscript
from app.statistics.reporting import fmt_df, fmt_p, fmt_stat

LABELS = {
    "tr": {
        "intro": "Giriş", "methods": "Yöntem", "results": "Bulgular",
        "exploratory": "Keşifsel Bulgular (Hipotez Üretici)",
        "discussion": "Tartışma", "limitations": "Sınırlılıklar", "references": "Kaynakça",
        "appendix": "Ek: Veri Temizleme Raporu",
        "appendix2": "Ek 2: Keşifsel Analiz Detayları",
        "table_results": "Tablo 1. Analiz Sonuçları Özeti",
        "table_clusters": "Tablo 2. Gizli Grup (Küme) Profilleri",
        "table_risk": "Tablo 3. Risk Skoru Yordayıcıları",
        "author": "Yazar Adı", "cols": ["Analiz", "Değişkenler", "İstatistik", "p", "Etki Büyüklüğü"],
        "cols_clusters": ["Grup", "n", "%", "Tanımlayıcı Değişkenler"],
        "cols_risk": ["Değişken", "Odds Oranı", "RF Önem Skoru"],
        "anomalies_line": "Sıra dışı vaka sayısı (Isolation Forest, kirlilik %{c:.0f}): {n}",
        "mi_line": "Gizli bilgi-teorik ilişki (MI): {a} – {b} (MI={mi:.3f}, r={r})",
        "skipped_line": "Atlanan analiz: {reason}",
    },
    "en": {
        "intro": "Introduction", "methods": "Method", "results": "Results",
        "exploratory": "Exploratory Findings (Hypothesis-Generating)",
        "discussion": "Discussion", "limitations": "Limitations", "references": "References",
        "appendix": "Appendix: Data Cleaning Report",
        "appendix2": "Appendix 2: Exploratory Analysis Details",
        "table_results": "Table 1. Summary of Analyses",
        "table_clusters": "Table 2. Hidden Cluster Profiles",
        "table_risk": "Table 3. Risk Score Predictors",
        "author": "Author Name", "cols": ["Analysis", "Variables", "Statistic", "p", "Effect Size"],
        "cols_clusters": ["Cluster", "n", "%", "Defining Variables"],
        "cols_risk": ["Variable", "Odds Ratio", "RF Importance"],
        "anomalies_line": "Flagged atypical cases (Isolation Forest, contamination {c:.0f}%): {n}",
        "mi_line": "Hidden information-theoretic relationship (MI): {a} – {b} (MI={mi:.3f}, r={r})",
        "skipped_line": "Skipped analysis: {reason}",
    },
}

TEST_NAMES = {
    "ttest_ind": {"tr": "Bağımsız t-testi", "en": "Independent t-test"},
    "welch_t": {"tr": "Welch t-testi", "en": "Welch's t-test"},
    "mannwhitney": {"tr": "Mann-Whitney U", "en": "Mann-Whitney U"},
    "anova": {"tr": "Tek yönlü ANOVA", "en": "One-way ANOVA"},
    "kruskal": {"tr": "Kruskal-Wallis H", "en": "Kruskal-Wallis H"},
    "chi2": {"tr": "Ki-kare", "en": "Chi-square"},
    "fisher": {"tr": "Fisher kesin testi", "en": "Fisher's exact test"},
    "pearson": {"tr": "Pearson korelasyonu", "en": "Pearson correlation"},
    "spearman": {"tr": "Spearman korelasyonu", "en": "Spearman correlation"},
    "linreg": {"tr": "Çoklu regresyon", "en": "Multiple regression"},
}


def _setup_styles(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    pf = style.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    pf.space_after = Pt(0)
    for section in doc.sections:
        section.top_margin = section.bottom_margin = Cm(2.5)
        section.left_margin = section.right_margin = Cm(2.5)


def _heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    if level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _body(doc: Document, text: str) -> None:
    for para in text.split("\n\n"):
        para = para.strip()
        if not para:
            continue
        p = doc.add_paragraph(para)
        p.paragraph_format.first_line_indent = Cm(1.25)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def add_title_page(doc: Document, title: str, lang: str) -> None:
    L = LABELS[lang]
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run(title or ("Başlıksız Çalışma" if lang == "tr" else "Untitled Study"))
    run.bold = True
    for line in (L["author"], ""):
        p = doc.add_paragraph(line)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()


def add_results_table(doc: Document, findings: list[Finding], lang: str) -> None:
    """Tablo 1: klasik hipotez testi sonuçları (doğrulayıcı istatistikler)."""
    L = LABELS[lang]
    valid = [f for f in findings if f.error is None]
    if not valid:
        return
    doc.add_paragraph()
    cap = doc.add_paragraph()
    cap.add_run(L["table_results"]).bold = True
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    for i, col in enumerate(L["cols"]):
        table.rows[0].cells[i].paragraphs[0].add_run(col).bold = True
    for f in valid:
        row = table.add_row().cells
        row[0].text = TEST_NAMES.get(f.planned.test_id, {}).get(lang, f.planned.test_id)
        vars_text = " × ".join(x for x in [f.planned.dv, f.planned.iv] if x) or ", ".join(f.planned.extra_vars)
        row[1].text = vars_text
        df_part = f"({fmt_df(f.df)})" if f.df is not None else ""
        row[2].text = f"{f.statistic_name}{df_part} = {fmt_stat(f.statistic)}"
        p_final = f.p_adjusted if f.p_adjusted is not None else f.p_value
        row[3].text = fmt_p(p_final) + (" *" if f.significant else "")
        row[4].text = f"{f.effect_size_name} = {fmt_stat(f.effect_size)}" if f.effect_size is not None else "—"


def add_exploratory_section(doc: Document, exploratory_text: str, discovery: DiscoveryReport | None, lang: str) -> None:
    """Keşifsel bulgular bölümü + küme/risk skoru tabloları — doğrulayıcı sonuçlardan ayrı."""
    L = LABELS[lang]
    if not exploratory_text:
        return
    _heading(doc, L["exploratory"])
    _body(doc, exploratory_text)

    if discovery and discovery.clustering and discovery.clustering.clusters:
        doc.add_paragraph()
        cap = doc.add_paragraph()
        cap.add_run(L["table_clusters"]).bold = True
        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        for i, col in enumerate(L["cols_clusters"]):
            table.rows[0].cells[i].paragraphs[0].add_run(col).bold = True
        for cl in discovery.clustering.clusters:
            row = table.add_row().cells
            row[0].text = str(cl.cluster_id)
            row[1].text = str(cl.size)
            row[2].text = f"{cl.share * 100:.0f}"
            row[3].text = ", ".join(v["name"] for v in cl.top_variables[:5])

    if discovery and discovery.risk_score and discovery.risk_score.predictors:
        doc.add_paragraph()
        cap = doc.add_paragraph()
        cap.add_run(L["table_risk"]).bold = True
        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        for i, col in enumerate(L["cols_risk"]):
            table.rows[0].cells[i].paragraphs[0].add_run(col).bold = True
        for p in discovery.risk_score.predictors:
            row = table.add_row().cells
            row[0].text = str(p["name"])
            row[1].text = fmt_stat(p["odds_ratio"])
            row[2].text = fmt_stat(p["rf_importance"], 3)


def add_cleaning_appendix(doc: Document, cleaning_report: CleaningReport | None, lang: str) -> None:
    if not cleaning_report:
        return
    L = LABELS[lang]
    doc.add_page_break()
    _heading(doc, L["appendix"])
    for action in cleaning_report.actions:
        doc.add_paragraph("• " + action)


def add_discovery_appendix(doc: Document, discovery: DiscoveryReport | None, lang: str) -> None:
    if not discovery or not (discovery.anomalies or discovery.mutual_info or discovery.skipped_reasons):
        return
    L = LABELS[lang]
    doc.add_page_break()
    _heading(doc, L["appendix2"])
    if discovery.anomalies:
        doc.add_paragraph(
            "• " + L["anomalies_line"].format(c=discovery.anomalies.contamination * 100, n=discovery.anomalies.n_flagged)
        )
    for p in discovery.mutual_info:
        if not p.hidden:
            continue
        r_txt = fmt_stat(p.correlation) if p.correlation is not None else "—"
        doc.add_paragraph("• " + L["mi_line"].format(a=p.var_a, b=p.var_b, mi=p.mi, r=r_txt))
    for reason in discovery.skipped_reasons:
        doc.add_paragraph("• " + L["skipped_line"].format(reason=reason))


def build_docx(
    manuscript: Manuscript,
    findings: list[Finding],
    cleaning_report: CleaningReport | None,
    discovery: DiscoveryReport | None,
    out_path: str | Path,
) -> Path:
    lang = manuscript.language
    L = LABELS[lang]
    doc = Document()
    _setup_styles(doc)

    add_title_page(doc, manuscript.title, lang)

    for key in ("intro", "methods", "results"):
        if manuscript.sections.get(key):
            _heading(doc, L[key])
            _body(doc, manuscript.sections[key])

    add_results_table(doc, findings, lang)
    add_exploratory_section(doc, manuscript.sections.get("exploratory", ""), discovery, lang)

    for key in ("discussion", "limitations"):
        if manuscript.sections.get(key):
            _heading(doc, L[key])
            _body(doc, manuscript.sections[key])

    # Kaynakça — asılı girinti, italik dergi adı
    if manuscript.references:
        doc.add_page_break()
        _heading(doc, L["references"])
        for ref in manuscript.references:
            p = doc.add_paragraph()
            pf = p.paragraph_format
            pf.left_indent = Cm(1.25)
            pf.first_line_indent = Cm(-1.25)
            for text, italic in format_reference_runs(ref):
                r = p.add_run(text)
                r.italic = italic

    add_cleaning_appendix(doc, cleaning_report, lang)
    add_discovery_appendix(doc, discovery, lang)

    out_path = Path(out_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return out_path
