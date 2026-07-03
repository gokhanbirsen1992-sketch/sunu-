"""Şablon modu yazımı ve Word çıktısı."""
from __future__ import annotations

from docx import Document

from app.llm import template_provider as tpl
from app.manuscript.docx_builder import build_docx
from app.models import CleaningReport, Finding, Manuscript, PlannedTest, Reference


def _finding(fid="F1", sig=True):
    return Finding(
        id=fid,
        planned=PlannedTest(test_id="ttest_ind", family="group", dv="kaygi", iv="cinsiyet",
                            rationale_tr="t-testi kullanıldı.", rationale_en="t-test was used."),
        statistic_name="t", statistic=2.31, df=58.0, p_value=0.024, p_adjusted=0.048,
        effect_size_name="d", effect_size=0.61, significant=sig,
        apa_tr="'Kaygı', 'Cinsiyet' gruplarına göre anlamlı fark göstermiştir (t(58) = 2.31, p = .024).",
        apa_en="'Anxiety' differed significantly across 'Gender' groups (t(58) = 2.31, p = .024).",
        keywords=["Kaygı Puanı", "Cinsiyet"],
        group_stats=[{"group": "Kadın", "n": 30, "mean": 56.0, "sd": 9.1, "median": 55.5}],
    )


def _refs(n=3):
    return [
        Reference(id=i, title=f"Study {i}", authors=[f"Author {i} Lastname{i}"],
                  year=2018 + i, journal="Journal of Tests", doi=f"10.1/x{i}")
        for i in range(1, n + 1)
    ]


def test_template_intro_and_discussion_cite_only_valid_ids():
    refs = _refs(3)
    findings = [_finding()]
    intro = tpl.draft_intro(findings, refs, "Kaygı ve doyum ilişkisi", "tr")
    disc = tpl.draft_discussion(findings, {"F1": refs}, "tr")
    from app.manuscript.citations import extract_marker_ids

    assert set(extract_marker_ids(intro)) <= {1, 2, 3}
    assert set(extract_marker_ids(disc)) <= {1, 2, 3}
    assert len(intro.split()) > 60
    assert "t(58)" in disc  # gerçek istatistik cümlesi taşınmış


def test_template_critiques_flag_low_refs():
    crits = tpl.template_critiques({"intro": "kısa", "discussion": "x"}, n_refs=2, lang="tr")
    assert any(c["requires_new_literature"] for c in crits)


def test_docx_builder_produces_valid_document(tmp_path):
    m = Manuscript(
        title="Deneme Makalesi", language="tr",
        sections={
            "intro": "Giriş metni (Lastname1, 2019).",
            "methods": "Yöntem metni.",
            "results": "Bulgular metni.",
            "discussion": "Tartışma metni (Lastname2, 2020).",
            "limitations": "Sınırlılıklar metni.",
        },
        references=_refs(2),
    )
    report = CleaningReport(rows_before=120, rows_after=115, actions=["Temizlik yapıldı."])
    out = build_docx(m, [_finding()], report, None, tmp_path / "out.docx")
    assert out.exists() and out.stat().st_size > 5000
    doc = Document(str(out))
    text = "\n".join(p.text for p in doc.paragraphs)
    for expected in ("Deneme Makalesi", "Giriş", "Yöntem", "Bulgular", "Tartışma", "Kaynakça"):
        assert expected in text
    assert len(doc.tables) == 1
    assert doc.tables[0].rows[1].cells[0].text == "Bağımsız t-testi"
