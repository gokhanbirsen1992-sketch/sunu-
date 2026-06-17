"""word2pdf paketi için testler.

Birim testleri her zaman çalışır; gerçek dönüştürme (entegrasyon) testleri
ilgili motor (LibreOffice ya da python-docx+reportlab) yoksa atlanır.
"""

from __future__ import annotations

from pathlib import Path

import pytest

from word2pdf import converter
from word2pdf.converter import (
    ConversionError,
    _resolve_output_path,
    convert_file,
    find_libreoffice,
    iter_word_files,
)

# --- Opsiyonel bağımlılıkların varlığı ---
try:
    import docx  # noqa: F401
    import reportlab  # noqa: F401

    HAS_PY_ENGINE = True
except ImportError:
    HAS_PY_ENGINE = False

HAS_LIBREOFFICE = find_libreoffice() is not None


def _libreoffice_functional() -> bool:
    """LibreOffice yalnızca kurulu değil, gerçekten dönüştürebiliyor mu?

    Bazı kurulumlarda çalıştırılabilir mevcuttur ama Writer bileşeni eksiktir;
    bu durumda dönüştürme başarısız olur. Küçük bir deneme dönüştürmesiyle
    fiilî yeteneği ölçeriz.
    """
    if not HAS_LIBREOFFICE or not HAS_PY_ENGINE:
        return False
    import tempfile

    with tempfile.TemporaryDirectory() as d:
        src = _make_sample_docx(Path(d) / "probe.docx")
        try:
            convert_file(src, Path(d) / "probe.pdf", engine="libreoffice")
        except ConversionError:
            return False
    return True


# --------------------------------------------------------------------------- #
# Yardımcılar
# --------------------------------------------------------------------------- #
def _make_sample_docx(path: Path) -> Path:
    """Türkçe karakterler, başlık, biçimlendirme ve tablo içeren bir .docx üretir."""
    import docx  # yerel import: HAS_PY_ENGINE False ise testler zaten atlanır

    doc = docx.Document()
    doc.add_heading("Şirket Faaliyet Raporu", level=1)
    doc.add_heading("Özet", level=2)
    p = doc.add_paragraph("Bu belge ")
    p.add_run("kalın").bold = True
    p.add_run(" ve ")
    p.add_run("italik").italic = True
    p.add_run(" metin içerir. Türkçe karakterler: çğıöşü ÇĞİÖŞÜ.")

    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Ürün"
    table.cell(0, 1).text = "Adet"
    table.cell(1, 0).text = "Çikolata"
    table.cell(1, 1).text = "42"

    doc.save(str(path))
    return path


def _assert_valid_pdf(path: Path) -> None:
    assert path.exists(), f"PDF üretilmedi: {path}"
    data = path.read_bytes()
    assert data.startswith(b"%PDF"), "Dosya geçerli bir PDF başlığıyla başlamıyor"
    assert len(data) > 200, "PDF beklenmedik şekilde küçük"


# Fiilî LibreOffice yeteneğini yardımcılar tanımlandıktan sonra ölç.
LIBREOFFICE_FUNCTIONAL = _libreoffice_functional()


# --------------------------------------------------------------------------- #
# Birim testleri (bağımsız)
# --------------------------------------------------------------------------- #
def test_resolve_output_default():
    out = _resolve_output_path(Path("/tmp/rapor.docx"), None)
    assert out == Path("/tmp/rapor.pdf")


def test_resolve_output_explicit_file():
    out = _resolve_output_path(Path("rapor.docx"), Path("ciktilar/sonuc.pdf"))
    assert out == Path("ciktilar/sonuc.pdf")


def test_resolve_output_directory_like(tmp_path):
    out_dir = tmp_path / "pdfler"
    out_dir.mkdir()
    out = _resolve_output_path(Path("a/rapor.docx"), out_dir)
    assert out == out_dir / "rapor.pdf"


def test_resolve_output_trailing_slash():
    out = _resolve_output_path(Path("rapor.docx"), Path("klasor/"))
    assert out == Path("klasor/rapor.pdf")


def test_missing_input_raises(tmp_path):
    with pytest.raises(ConversionError, match="bulunamadı"):
        convert_file(tmp_path / "yok.docx")


def test_unsupported_extension_raises(tmp_path):
    bad = tmp_path / "not.txt"
    bad.write_text("merhaba")
    with pytest.raises(ConversionError, match="Desteklenmeyen uzantı"):
        convert_file(bad)


def test_unknown_engine_raises(tmp_path):
    src = tmp_path / "a.docx"
    src.write_bytes(b"PK\x03\x04")  # geçerlilik aranmadan önce motor kontrolü için
    with pytest.raises(ConversionError, match="Bilinmeyen motor"):
        convert_file(src, engine="sihirli")


def test_find_libreoffice_type():
    result = find_libreoffice()
    assert result is None or isinstance(result, str)


@pytest.mark.skipif(not HAS_PY_ENGINE, reason="python-docx/reportlab kurulu değil")
def test_iter_word_files(tmp_path):
    _make_sample_docx(tmp_path / "a.docx")
    _make_sample_docx(tmp_path / "b.docx")
    (tmp_path / "not.txt").write_text("x")
    (tmp_path / "~$gecici.docx").write_text("kilit")  # Office kilit dosyası
    sub = tmp_path / "alt"
    sub.mkdir()
    _make_sample_docx(sub / "c.docx")

    flat = list(iter_word_files(tmp_path, recursive=False))
    assert {p.name for p in flat} == {"a.docx", "b.docx"}

    deep = list(iter_word_files(tmp_path, recursive=True))
    assert {p.name for p in deep} == {"a.docx", "b.docx", "c.docx"}


# --------------------------------------------------------------------------- #
# Entegrasyon testleri (motor gerektirir)
# --------------------------------------------------------------------------- #
@pytest.mark.skipif(not HAS_PY_ENGINE, reason="python-docx/reportlab kurulu değil")
def test_convert_python_engine(tmp_path):
    src = _make_sample_docx(tmp_path / "rapor.docx")
    out = convert_file(src, engine="python")
    assert out == tmp_path / "rapor.pdf"
    _assert_valid_pdf(out)


@pytest.mark.skipif(
    not LIBREOFFICE_FUNCTIONAL,
    reason="Çalışan bir LibreOffice (Writer bileşeniyle) yok",
)
def test_convert_libreoffice_engine(tmp_path):
    src = _make_sample_docx(tmp_path / "rapor.docx")
    out = convert_file(src, tmp_path / "cikti.pdf", engine="libreoffice")
    assert out == tmp_path / "cikti.pdf"
    _assert_valid_pdf(out)


@pytest.mark.skipif(not HAS_PY_ENGINE, reason="python-docx/reportlab kurulu değil")
def test_overwrite_guard(tmp_path):
    src = _make_sample_docx(tmp_path / "rapor.docx")
    convert_file(src, engine="python")
    # İkinci kez üzerine yazmadan dönüştürmek hata vermeli.
    with pytest.raises(ConversionError, match="zaten var"):
        convert_file(src, engine="python")
    # overwrite=True ile sorunsuz çalışmalı.
    out = convert_file(src, engine="python", overwrite=True)
    _assert_valid_pdf(out)


@pytest.mark.skipif(not HAS_PY_ENGINE, reason="python-docx/reportlab kurulu değil")
def test_output_to_directory(tmp_path):
    src = _make_sample_docx(tmp_path / "rapor.docx")
    out_dir = tmp_path / "pdfler"
    out_dir.mkdir()
    out = convert_file(src, out_dir, engine="python")
    assert out == out_dir / "rapor.pdf"
    _assert_valid_pdf(out)
