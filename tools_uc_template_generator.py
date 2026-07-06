# -*- coding: utf-8 -*-
"""Ülseratif Kolit Poliklinik Takip Şablonu (.docx) üreticisi."""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ACCENT = RGBColor(0x1F, 0x4E, 0x79)
LIGHT = RGBColor(0x5B, 0x6B, 0x7A)


def set_cell_bg(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hexcolor)
    tcPr.append(shd)


def h(doc, text, size=13, color=ACCENT, space_before=10, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size)
    r.font.color.rgb = color
    return p


def field(doc, label, hint=""):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("• " + label + ": ")
    r.bold = True
    r.font.size = Pt(10)
    if hint:
        rh = p.add_run("[ " + hint + " ]")
        rh.italic = True
        rh.font.size = Pt(9)
        rh.font.color.rgb = LIGHT
    return p


def note(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(8.5)
    r.font.color.rgb = LIGHT
    return p


def checkbox_line(doc, options):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    for i, opt in enumerate(options):
        r = p.add_run("☐ " + opt)
        r.font.size = Pt(10)
        if i < len(options) - 1:
            sp = p.add_run("    ")
            sp.font.size = Pt(10)
    return p


def make_table(doc, header, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(header))
    t.style = 'Light Grid Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, htext in enumerate(header):
        hdr[i].text = ""
        pr = hdr[i].paragraphs[0].add_run(htext)
        pr.bold = True
        pr.font.size = Pt(9)
        pr.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_bg(hdr[i], '1F4E79')
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            rr = cells[i].paragraphs[0].add_run(val)
            rr.font.size = Pt(9)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    return t


doc = Document()
# Sayfa kenar boşlukları
for s in doc.sections:
    s.top_margin = Cm(1.6)
    s.bottom_margin = Cm(1.6)
    s.left_margin = Cm(1.8)
    s.right_margin = Cm(1.8)

normal = doc.styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(10)

# ---- BAŞLIK ----
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = title.add_run("ÜLSERATİF KOLİT — POLİKLİNİK TAKİP FORMU")
tr.bold = True
tr.font.size = Pt(16)
tr.font.color.rgb = ACCENT

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = sub.add_run("Gastroenteroloji / İç Hastalıkları Polikliniği — İBH İzlem Formu")
sr.font.size = Pt(10)
sr.italic = True
sr.font.color.rgb = LIGHT

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
mr = meta.add_run("Literatür temeli: ECCO Kılavuzları (Tanı/İzlem 2022, Tıbbi Tedavi 2022), "
                  "STRIDE-II Treat-to-Target (2021), Mayo & UCEIS aktivite skorları")
mr.font.size = Pt(8)
mr.font.color.rgb = LIGHT

# ---- 1. HASTA KİMLİK / ZİYARET ----
h(doc, "1. HASTA KİMLİK VE ZİYARET BİLGİLERİ")
row = doc.add_table(rows=1, cols=2)
c = row.rows[0].cells
for cell, items in zip(c, [
    ["Ad-Soyad / Protokol No:", "Doğum tarihi / Yaş:", "Cinsiyet:", "Telefon:"],
    ["Ziyaret tarihi:", "Ziyaret tipi:  ☐ Rutin  ☐ Alevlenme  ☐ Kontrol", "Takip eden hekim:", "Bir sonraki randevu:"]]):
    for it in items:
        pp = cell.add_paragraph()
        pp.paragraph_format.space_after = Pt(3)
        run = pp.add_run(it)
        run.font.size = Pt(10)
    # ilk boş paragrafı sil
    cell.paragraphs[0]._p.getparent().remove(cell.paragraphs[0]._p) if cell.paragraphs[0].text == "" and len(cell.paragraphs) > 1 else None

# ---- 2. TANI VE HASTALIK ÖZELLİKLERİ ----
h(doc, "2. TANI VE HASTALIK ÖZELLİKLERİ (Montreal Sınıflaması)")
field(doc, "Tanı tarihi / Hastalık süresi", "yıl")
field(doc, "Tanı yöntemi", "kolonoskopi + histopatoloji")
doc.add_paragraph().add_run("Yaygınlık (Montreal E):").bold = True
checkbox_line(doc, ["E1 – Proktit (rektum)", "E2 – Sol kolon (splenik fleksuraya dek)", "E3 – Yaygın / Pankolit"])
note(doc, "E: hastalığın en geniş yayıldığı anatomik sınır (endoskopik). Yaygınlık; gözetim, tedavi yolu ve kolektomi riskini belirler (ECCO 2022).")

doc.add_paragraph().add_run("Şiddet (Montreal S — güncel/en yüksek):").bold = True
checkbox_line(doc, ["S0 – Klinik remisyon", "S1 – Hafif", "S2 – Orta", "S3 – Ağır"])

field(doc, "Ekstraintestinal tutulum", "artrit/artralji, üveit/episklerit, eritema nodozum, piyoderma gangrenozum, PSK, VTE öyküsü")
field(doc, "Aile öyküsü (İBH)", "")
field(doc, "Sigara durumu", "hiç / bırakmış / aktif")
field(doc, "Önceki komplikasyonlar", "ağır atak/hastane yatışı, toksik megakolon, displazi, cerrahi")

# ---- 3. GÜNCEL TEDAVİ ----
h(doc, "3. GÜNCEL TEDAVİ VE UYUM")
make_table(doc, ["İlaç", "Doz / Yol", "Başlangıç", "Uyum (İyi/Orta/Kötü)", "Yan etki"],
           [["", "", "", "", ""] for _ in range(4)])
note(doc, "5-ASA (oral ± topikal), topikal steroid, sistemik steroid, tiyopürin (AZA/6-MP), anti-TNF (infliksimab, adalimumab, golimumab), "
          "vedolizumab, ustekinumab, tofasitinib/upadasitinib, ozanimod. Steroid; idame tedavi DEĞİLDİR — steroid bağımlılığı/direnci not edilmeli.")
field(doc, "Steroide maruziyet (son 12 ay)", "kür sayısı / kümülatif süre")

# ---- 4. SEMPTOM DEĞERLENDİRME (pMayo/SCCAI) ----
h(doc, "4. SEMPTOM DEĞERLENDİRME — Kısmi Mayo (pMayo) Skoru")
make_table(doc,
    ["Parametre", "0", "1", "2", "3"],
    [
     ["Dışkılama sıklığı", "Normal", "1–2/gün fazla", "3–4/gün fazla", "≥5/gün fazla"],
     ["Rektal kanama", "Yok", "Çizgi şeklinde (<%50 zaman)", "Belirgin (≥%50 zaman)", "Yalnız kan"],
     ["Hekim genel değerlendirmesi", "Normal", "Hafif", "Orta", "Ağır"],
    ],
    widths=[5.5, 2.6, 3.5, 3.5, 2.8])
field(doc, "pMayo toplam (0–9)", "")
note(doc, "Yorum: Remisyon pMayo ≤2 ve hiçbir alt skor >1; ayrıca rektal kanama alt skoru = 0 hedeflenir (STRIDE-II klinik hedef). "
          "Yanıt: başlangıca göre ≥2 puan ve ≥%30 düşüş. Alternatif: SCCAI ≤2 remisyon.")

field(doc, "Ek semptomlar", "karın ağrısı, aciliyet/urgency, gece dışkılama, tenezm, kilo kaybı, ateş")
field(doc, "Yaşam kalitesi / işe-okula devamsızlık", "IBD-Disk veya sIBDQ opsiyonel")
field(doc, "Dışkılama sıklığı (gündüz/gece)", "")

# ---- 5. FİZİK MUAYENE ----
h(doc, "5. FİZİK MUAYENE")
field(doc, "Genel durum / Vital", "Ateş, NB, TA, KH")
field(doc, "Kilo / Boy / VKİ", "kilo trendi önemli")
field(doc, "Batın muayenesi", "hassasiyet, distansiyon, rebound")
field(doc, "Perianal / rektal muayene", "")
field(doc, "Ekstraintestinal bulgular", "eklem, göz, cilt, oral aft")

# ---- 6. LABORATUVAR ----
h(doc, "6. LABORATUVAR VE BİYOBELİRTEÇLER")
make_table(doc, ["Test", "Sonuç", "Referans / Hedef", "Not"],
    [
     ["Hemogram (Hb, Lökosit, Trombosit)", "", "Hb: anemi taraması", ""],
     ["CRP", "", "Normalleşme hedef", "Nonspesifik; tek başına yetersiz"],
     ["Sedimantasyon (ESR)", "", "", ""],
     ["Fekal Kalprotektin", "", "< 150 µg/g (hedef; ideal <100–250)", "Mukozal aktivitenin en iyi non-invaziv göstergesi"],
     ["Albümin / Total protein", "", "Ağır atakta düşük", ""],
     ["Ferritin / Demir / TSAT / B12–Folat", "", "Anemi etiyolojisi", ""],
     ["Elektrolit / Böbrek / KCFT", "", "İlaç izlemi", ""],
     ["Gaita mikrobiyolojisi + C. difficile toksini", "", "Alevlenmede zorunlu", "Enfeksiyonu dışla"],
    ],
    widths=[6.5, 2.6, 4.4, 4.4])
note(doc, "Fekal kalprotektin STRIDE-II'de ara (formatif) hedeftir; kalıcı yükseklik endoskopik değerlendirmeyi tetikler. "
          "Tiyopürin öncesi TPMT/NUDT15; biyolojik/JAK öncesi tarama testleri (aşağıda).")

# ---- 7. ENDOSKOPİK / GÖRÜNTÜLEME ----
h(doc, "7. ENDOSKOPİK AKTİVİTE VE GÖRÜNTÜLEME")
field(doc, "Son kolonoskopi/sigmoidoskopi tarihi", "")
doc.add_paragraph().add_run("Mayo Endoskopik Alt Skoru (MES):").bold = True
checkbox_line(doc, ["0 – Normal / inaktif", "1 – Hafif (eritem, azalmış vasküler ağ)",
                    "2 – Orta (belirgin eritem, erozyon)", "3 – Ağır (spontan kanama, ülser)"])
field(doc, "UCEIS (0–8)", "vasküler patern + kanama + erozyon/ülser")
note(doc, "Endoskopik iyileşme (mukozal iyileşme) = MES ≤1 (ideal MES 0). STRIDE-II uzun dönem hedefi. "
          "UCEIS daha ayrıntılı ve doğrulanmış; MES ile birlikte kullanılabilir. Histolojik remisyon opsiyonel ileri hedeftir.")
field(doc, "Histoloji", "Nancy/Robarts indeksi — bazal plazmositoz, nötrofilik infiltrasyon")
field(doc, "Görüntüleme (gerekirse)", "ADBG/BT — ağır atakta toksik megakolon şüphesi")

# ---- 8. HASTALIK AKTİVİTE SINIFLAMASI ----
h(doc, "8. HASTALIK AKTİVİTESİ — Truelove-Witts (Ağır Atak Değerlendirmesi)")
make_table(doc, ["Kriter", "Hafif", "Ağır (ATAK)"],
    [
     ["Kanlı dışkılama / gün", "< 4", "≥ 6"],
     ["Nabız", "< 90/dk", "> 90/dk"],
     ["Ateş", "< 37.8°C", "> 37.8°C"],
     ["Hemoglobin", "Normal", "< 10.5 g/dL"],
     ["ESR / CRP", "≤ 30 mm/s", "> 30 mm/s"],
    ],
    widths=[7.0, 5.4, 5.4])
note(doc, "≥6 kanlı dışkılama + en az 1 sistemik bulgu (nabız/ateş/Hb/ESR) = AĞIR ATAK (Akut Ağır Ülseratif Kolit — ASUC). "
          "Hastaneye yatış, IV steroid, VTE profilaksisi, günlük izlem; 3. günde Oxford/Travis kriterleri ile kurtarma tedavisi kararı.")

doc.add_paragraph().add_run("Genel değerlendirme:").bold = True
checkbox_line(doc, ["Klinik remisyon", "Hafif aktivite", "Orta aktivite", "Ağır atak (ASUC)"])

# ---- 9. TREAT-TO-TARGET (STRIDE-II) ----
h(doc, "9. TEDAVİ HEDEFLERİ — Treat-to-Target (STRIDE-II 2021)")
make_table(doc, ["Zaman", "Hedef", "Ulaşıldı mı?"],
    [
     ["Kısa dönem", "Klinik yanıt (semptomlarda düzelme)", "☐ Evet  ☐ Hayır"],
     ["Orta dönem", "Klinik remisyon (pMayo ≤2) + CRP normal + FKP düşüş", "☐ Evet  ☐ Hayır"],
     ["Uzun dönem", "Endoskopik iyileşme (MES ≤1) + normal yaşam kalitesi + kalıcı steroidsiz remisyon", "☐ Evet  ☐ Hayır"],
    ],
    widths=[3.5, 10.5, 3.8])
note(doc, "Hedefe ulaşılmadıysa tedavi optimizasyonu/değişikliği düşün (doz artışı, TDM ile ilaç düzeyi/antikor, sınıf değişimi). "
          "Formatif hedefler: CRP ve fekal kalprotektin normalizasyonu.")

# ---- 10. GÖZETİM / TARAMA / KORUYUCU ----
h(doc, "10. GÖZETİM, TARAMA VE KORUYUCU SAĞLIK")
make_table(doc, ["Alan", "Öneri", "Son / Planlanan"],
    [
     ["KRK gözetim kolonoskopisi", "Yaygın/sol kolon tutulumunda tanıdan 8 yıl sonra başla; risk katmanına göre 1–5 yıl (kromoendoskopi/HD)", ""],
     ["PSK varlığı", "Tanı anından itibaren yıllık kolonoskopi", ""],
     ["Kemik sağlığı (DEXA)", "Uzun/tekrarlayan steroid kullanımında; Ca/D vitamini", ""],
     ["Aşılama", "İnaktive: influenza, pnömokok, Hep B, HPV, COVID-19; canlı aşılar immünsupresyon öncesi", ""],
     ["Latent enfeksiyon taraması", "Biyolojik/JAK öncesi TB (IGRA/PPD, akciğer grafisi), HBV, HCV, HIV, VZV", ""],
     ["Cilt / dermatolojik", "Tiyopürin: NMSC riski; JAK/anti-TNF: cilt izlemi", ""],
     ["Ruhsal sağlık", "Anksiyete/depresyon taraması", ""],
     ["Üreme sağlığı", "Gebelik planı, gebelikte remisyon hedefi, ilaç uyumluluğu", ""],
    ],
    widths=[4.2, 9.6, 4.0])

# ---- 11. DEĞERLENDİRME VE PLAN ----
h(doc, "11. DEĞERLENDİRME VE PLAN")
field(doc, "Klinik değerlendirme özeti", "aktivite, hedefe ulaşım, komplikasyon")
doc.add_paragraph().add_run("Tedavi kararı:").bold = True
checkbox_line(doc, ["Mevcut tedaviye devam", "Doz optimizasyonu", "Basamak yükseltme (step-up)",
                    "Sınıf değişimi", "Topikal ekleme", "Cerrahi konsültasyon"])
field(doc, "Yeni/değişen ilaç ve gerekçe", "")
field(doc, "TDM / ilaç düzeyi istemi", "anti-TNF düzeyi + anti-drug antikor")
field(doc, "Beslenme / demir replasmanı", "IV demir tercih")
field(doc, "Hasta eğitimi verildi", "hastalık, ilaç uyumu, alevlenme belirtileri, sigara bırakma, aşı")
field(doc, "Kırmızı bayrak uyarıları anlatıldı", "artan kanama, ateş, şiddetli ağrı, kilo kaybı → acil başvuru")
field(doc, "Bir sonraki değerlendirme planı", "poliklinik / lab / FKP / endoskopi tarihi")

# İmza
doc.add_paragraph()
sig = doc.add_paragraph()
sr = sig.add_run("Hekim adı-soyadı / İmza: ______________________________        Tarih: ____/____/________")
sr.font.size = Pt(10)

# ---- KAYNAKLAR ----
h(doc, "Kaynaklar / Referanslar", size=11, space_before=14)
refs = [
    "Raine T, et al. ECCO Guidelines on Therapeutics in Ulcerative Colitis: Medical Treatment. J Crohns Colitis. 2022;16(1):2–17.",
    "Gordon H, et al. ECCO Guidelines on Therapeutics in UC: Surgical Treatment / Diagnosis & Monitoring (ECCO 2022–2023).",
    "Turner D, et al. STRIDE-II: An Update on the Selecting Therapeutic Targets in IBD (STRIDE) Initiative. Gastroenterology. 2021;160(5):1570–1583.",
    "Schroeder KW, et al. Mayo skoru. N Engl J Med. 1987;317:1625–1629.",
    "Travis SPL, et al. UCEIS — Ulcerative Colitis Endoscopic Index of Severity. Gut. 2012 & Gastroenterology 2013.",
    "Truelove SC, Witts LJ. Cortisone in ulcerative colitis. BMJ. 1955;2:1041–1048.",
    "Lamb CA, et al. British Society of Gastroenterology (BSG) IBD guidelines. Gut. 2019;68(Suppl 3).",
]
for r in refs:
    p = doc.add_paragraph(style='List Number')
    run = p.add_run(r)
    run.font.size = Pt(8.5)

disc = doc.add_paragraph()
dr = disc.add_run("Not: Bu form klinik karar desteği amaçlıdır; kurumsal protokoller ve güncel kılavuzlarla birlikte kullanılmalı, "
                  "hekimin klinik yargısının yerine geçmez. Eşik değerler kaynaklara göre değişebilir.")
dr.italic = True
dr.font.size = Pt(8)
dr.font.color.rgb = LIGHT

out = "/home/user/sunu-/Ulseratif_Kolit_Poliklinik_Takip_Formu.docx"
doc.save(out)
print("Kaydedildi:", out)
