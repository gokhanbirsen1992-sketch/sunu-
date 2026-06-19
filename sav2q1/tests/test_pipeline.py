"""Uçtan uca: motor → ledger → docx. M0 yürüyen iskeletini bütün olarak doğrular."""

import json

from docx import Document

from sav2q1.docx.assemble import build_docx


def test_ledger_built(demo_run):
    L = json.loads((demo_run / "results_ledger.json").read_text(encoding="utf-8"))
    r = L["results"][0]
    assert r["test"] == "mann_whitney_u"
    # Q1 şartı: etki büyüklüğü + %95 GA mevcut
    assert r["effect"]["name"] == "rank-biserial r"
    assert len(r["effect"]["ci"]) == 2
    # seçim gerekçesi kayıtlı (Yöntem dürüstlüğü)
    assert "Mann-Whitney" in r["reason"]
    # number_index binding-bazlı doğrulama için id'lere göre
    assert "R1" in L["number_index"]
    assert L["global_index"], "örneklem sayıları global_index'te olmalı"


def test_docx_built(demo_run):
    out = build_docx(str(demo_run))
    d = Document(out)
    texts = [p.text for p in d.paragraphs]
    assert any("Kaynaklar" in t for t in texts)
    assert any("Kunzler" in t for t in texts)          # GERÇEK atıf gömülü
    assert any(t.startswith("Öz") or t == "Öz" for t in texts)
    assert len(d.tables) == 1                            # Tablo 1
    # ICMJE beyanları
    assert any("Etik Kurul" in t for t in texts)
