"""Veri dosyasından tek seferde tam istatistik + keşifsel örüntü/risk skoru raporu üretir.

PaperForge'un web pipeline'ından (literatür taraması, makale yazımı, LLM) tamamen
bağımsızdır — hiçbir ağ çağrısı veya yapay zekâ kullanmaz; tüm sayılar
SciPy/statsmodels/scikit-learn ile deterministik hesaplanır.
"""
from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.manuscript.docx_builder import (
    LABELS,
    _body,
    _heading,
    _setup_styles,
    add_cleaning_appendix,
    add_discovery_appendix,
    add_exploratory_section,
    add_results_table,
)
from app.models import CleaningReport, DiscoveryReport, Finding, JobConfig, PlannedTest, VariableInfo
from app.pipeline.stages.s4_writing import build_exploratory, build_methods, build_results
from app.statistics.cleaning import apply_value_labels, clean
from app.statistics.decision import plan_tests
from app.statistics.discovery import run_discovery
from app.statistics.loader import load_dataset
from app.statistics.tests_runner import run_tests
from app.statistics.vartypes import infer_types

VAR_SUMMARY_HEADERS = {
    "tr": ["Değişken", "Etiket", "Tür", "Rol", "Eksik", "Benzersiz"],
    "en": ["Variable", "Label", "Kind", "Role", "Missing", "Unique"],
}


@dataclass
class AnalysisResult:
    variables: list[VariableInfo]
    plans: list[PlannedTest]
    findings: list[Finding]
    cleaning_report: CleaningReport
    discovery: DiscoveryReport
    n_rows_before: int
    n_rows_after: int
    out_path: Path


def analyze(
    input_path: str | Path,
    out_path: str | Path | None = None,
    dv: str | None = None,
    lang: str = "tr",
    alpha: float = 0.05,
    p_adjust: str = "holm",
) -> AnalysisResult:
    """Veriyi yükler, temizler, klasik testleri ve keşifsel analizi çalıştırır, Word raporu üretir."""
    input_path = Path(input_path)
    df, meta = load_dataset(input_path)
    variables = infer_types(df, meta)

    if dv:
        names = [v.name for v in variables]
        if dv not in names:
            raise ValueError(f"Değişken bulunamadı: '{dv}'. Mevcut değişkenler: {', '.join(names)}")
        for v in variables:
            if v.name == dv:
                v.role = "dv"

    n_rows_before = len(df)
    df_clean, cleaning_report = clean(df, variables)
    df_clean = apply_value_labels(df_clean, variables)

    config = JobConfig(language=lang, alpha=alpha, p_adjust=p_adjust)
    plans, notes = plan_tests(df_clean, variables)
    findings = run_tests(df_clean, plans, variables, config)
    discovery = run_discovery(df_clean, variables, config, findings)

    out_path = Path(out_path) if out_path else input_path.with_name(input_path.stem + "_rapor.docx")
    _build_report_docx(
        out_path, input_path.name, variables, cleaning_report, plans, findings, discovery, config, notes,
    )

    return AnalysisResult(
        variables=variables, plans=plans, findings=findings, cleaning_report=cleaning_report,
        discovery=discovery, n_rows_before=n_rows_before, n_rows_after=cleaning_report.rows_after,
        out_path=out_path,
    )


def _build_report_docx(
    out_path: Path,
    source_name: str,
    variables: list[VariableInfo],
    cleaning_report: CleaningReport,
    plans: list[PlannedTest],
    findings: list[Finding],
    discovery: DiscoveryReport,
    config: JobConfig,
    notes: list[str],
) -> None:
    lang = config.language
    L = LABELS[lang]
    doc = Document()
    _setup_styles(doc)

    title = ("İstatistik Analiz Raporu — " if lang == "tr" else "Statistical Analysis Report — ") + source_name
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.add_run(title).bold = True
    doc.add_page_break()

    _heading(doc, "Değişken Özeti" if lang == "tr" else "Variable Summary")
    headers = VAR_SUMMARY_HEADERS[lang]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, col in enumerate(headers):
        table.rows[0].cells[i].paragraphs[0].add_run(col).bold = True
    for v in variables:
        row = table.add_row().cells
        row[0].text = v.name
        row[1].text = v.label or ""
        row[2].text = v.kind
        row[3].text = v.role
        row[4].text = str(v.n_missing)
        row[5].text = str(v.n_unique)

    _heading(doc, L["methods"])
    _body(doc, build_methods(cleaning_report, plans, config))
    if notes:
        _body(doc, " ".join(notes))

    _heading(doc, L["results"])
    _body(doc, build_results(findings, lang))
    add_results_table(doc, findings, lang)

    add_exploratory_section(doc, build_exploratory(discovery, lang), discovery, lang)

    add_cleaning_appendix(doc, cleaning_report, lang)
    add_discovery_appendix(doc, discovery, lang)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
