"""Veri Kâşifi — Word rapor üretici. Tamamen bağımsız (python-docx dışında dış bağımlılık yok)."""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.shared import Cm, Pt

from analysis import ColumnInfo, DiscoveryResult, TestResult


def _setup_styles(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    for section in doc.sections:
        section.top_margin = section.bottom_margin = Cm(2.0)
        section.left_margin = section.right_margin = Cm(2.0)


def _heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)


def build_report(
    out_path: str | Path,
    source_name: str,
    columns: list[ColumnInfo],
    n_rows: int,
    tests: list[TestResult],
    tests_truncated: bool,
    discovery: DiscoveryResult,
) -> Path:
    doc = Document()
    _setup_styles(doc)

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.add_run(f"Veri Kâşifi — Analiz Raporu\n{source_name}").bold = True
    doc.add_paragraph()

    _heading(doc, "Değişken Özeti")
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    for i, col in enumerate(["Değişken", "Tür", "Eksik", "Benzersiz"]):
        table.rows[0].cells[i].paragraphs[0].add_run(col).bold = True
    for c in columns:
        row = table.add_row().cells
        row[0].text = c.name
        row[1].text = "Sayısal" if c.kind == "numeric" else "Kategorik"
        row[2].text = str(c.n_missing)
        row[3].text = str(c.n_unique)
    doc.add_paragraph(f"Toplam {n_rows} satır, {len(columns)} analiz edilebilir değişken.")

    doc.add_paragraph()
    _heading(doc, "Klasik İstatistik Testleri")
    sig = [t for t in tests if t.significant]
    doc.add_paragraph(f"{len(tests)} test çalıştırıldı, {len(sig)} tanesi istatistiksel olarak anlamlı (α=0.05).")
    if tests_truncated:
        doc.add_paragraph("Not: çok sayıda değişken kombinasyonu bulunduğundan yalnızca ilk sonuçlar listelenmiştir.")
    if tests:
        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        for i, col in enumerate(["Değişkenler", "İstatistik / p", "Yorum"]):
            table.rows[0].cells[i].paragraphs[0].add_run(col).bold = True
        for t in sorted(tests, key=lambda t: t.p_value)[:40]:
            row = table.add_row().cells
            row[0].text = " × ".join(t.variables)
            row[1].text = f"stat={t.statistic:.2f}, p={t.p_value:.3f}" + (" *" if t.significant else "")
            row[2].text = t.description

    doc.add_paragraph()
    _heading(doc, "Keşifsel Bulgular (Hipotez Üretici)")
    doc.add_paragraph(
        "Aşağıdaki bulgular klasik anlamlılık testlerinden farklıdır — kümeleme, anomali tespiti "
        "ve makine öğrenmesi tabanlı örüntü keşfine dayanır; doğrulayıcı değil hipotez üreticidir."
    )

    if discovery.clustering_k:
        doc.add_paragraph(
            f"K-Means kümeleme ile veride {discovery.clustering_k} gizli grup tespit edildi "
            f"(silhouette skoru = {discovery.clustering_silhouette:.2f})."
        )
        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        for i, col in enumerate(["Grup", "n", "%", "Tanımlayıcı Değişkenler"]):
            table.rows[0].cells[i].paragraphs[0].add_run(col).bold = True
        for cl in discovery.clusters:
            row = table.add_row().cells
            row[0].text = str(cl.cluster_id)
            row[1].text = str(cl.size)
            row[2].text = f"{cl.share * 100:.0f}"
            row[3].text = ", ".join(cl.top_variables)

    if discovery.n_anomalies is not None:
        doc.add_paragraph(
            f"Isolation Forest ile çok-değişkenli aykırı değer taramasında "
            f"{discovery.n_anomalies} sıra dışı vaka işaretlendi (kirlilik oranı %{discovery.anomaly_contamination * 100:.0f})."
        )

    if discovery.hidden_relationships:
        doc.add_paragraph(
            f"Karşılıklı bilgi (mutual information) analizi, klasik korelasyonun zayıf gösterdiği "
            f"ancak bilgi akışı taşıyan {len(discovery.hidden_relationships)} değişken çifti ortaya çıkardı:"
        )
        for p in discovery.hidden_relationships:
            r_txt = f"{p['r']:.2f}" if p["r"] is not None else "—"
            doc.add_paragraph(f"• {p['a']} – {p['b']} (MI={p['mi']:.3f}, r={r_txt})")

    if discovery.risk_score:
        r = discovery.risk_score
        auc = r["auc_logreg"] if r["auc_logreg"] is not None else r["auc_rf"]
        auc_txt = f"AUC={auc:.2f}" if auc is not None else "AUC hesaplanamadı"
        doc.add_paragraph(
            f"'{r['target']}' için lojistik regresyon + Random Forest ile risk skoru modeli kuruldu "
            f"(n={r['n']}, {auc_txt}). Pozitif sınıf: '{r['positive_class']}'."
        )
        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        for i, col in enumerate(["Değişken", "Önem Skoru"]):
            table.rows[0].cells[i].paragraphs[0].add_run(col).bold = True
        for p in r["predictors"]:
            row = table.add_row().cells
            row[0].text = p["name"]
            row[1].text = f"{p['importance']:.3f}"
        doc.add_paragraph("Bu skor, klinik/karar amaçlı kullanılmadan önce bağımsız bir örneklemde doğrulanmalıdır.")

    if discovery.notes:
        doc.add_paragraph()
        doc.add_paragraph("Atlanan analizler:").runs[0].italic = True
        for note in discovery.notes:
            doc.add_paragraph(f"• {note}")

    out_path = Path(out_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return out_path
